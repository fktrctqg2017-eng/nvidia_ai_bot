"""Обработка архивов, присланных пользователем: безопасная распаковка,
классификация содержимого (изображения / текст-код) и подготовка данных для
мультимодального запроса к LLM (Vision API).

Поддерживаемые форматы архивов: .zip, .tar, .tar.gz/.tgz, .tar.bz2/.tbz2,
.tar.xz/.txz, а также одиночные .gz файлы (не tar, просто сжатый файл).

Меры защиты при распаковке (важно, так как архивы — частый вектор атак),
одинаково применяются ко всем форматам:
- Path traversal / "Zip Slip" / "Tar Slip": проверяем, что итоговый путь
  каждого файла остаётся внутри целевой директории (запрещаем
  "../../etc/passwd", абсолютные пути и символические/жёсткие ссылки,
  указывающие за пределы директории).
- "Архивная бомба" (zip/tar bomb): ограничиваем суммарный распакованный
  размер, количество файлов и подозрительно высокую степень сжатия отдельных
  файлов, прерывая распаковку при превышении лимитов.
- Отдельная временная директория на каждый архив, гарантированно удаляется
  после обработки (см. `extract_archive_safely` как контекстный менеджер).
"""

from __future__ import annotations

import base64
import gzip
import mimetypes
import tarfile
import zipfile
from contextlib import contextmanager
from dataclasses import dataclass, field
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterator, Optional

# Расширения "документов" со СПЕЦИАЛЬНЫМ разбором (не просто текст построчно) —
# извлекаются через специализированные библиотеки (pypdf/python-docx/openpyxl).
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}
DOCUMENT_EXTENSIONS = PDF_EXTENSIONS | DOCX_EXTENSIONS | XLSX_EXTENSIONS
MAX_DOCUMENT_SIZE_BYTES = 20 * 1024 * 1024  # ограничение самого Telegram Bot API

# Сколько символов извлечённого из документа текста показываем модели/пользователю
# за раз (аналогично MAX_TEXT_FILE_CHARS_IN_PROMPT/STORAGE ниже для архивов).
MAX_DOCUMENT_CHARS_IN_PROMPT = 12000
MAX_DOCUMENT_CHARS_STORAGE = 300_000
MAX_XLSX_ROWS_PER_SHEET = 200   # защита от гигантских таблиц — обрезаем построчно
MAX_XLSX_SHEETS = 20

# ------------------------------------------------------------------ лимиты безопасности

MAX_TOTAL_UNCOMPRESSED_BYTES = 100 * 1024 * 1024  # 100 MB суммарно
MAX_FILES_IN_ARCHIVE = 200
MAX_SINGLE_FILE_BYTES = 20 * 1024 * 1024           # 20 MB на файл
MAX_COMPRESSION_RATIO = 100                        # защита от архивных бомб (высокая степень сжатия)

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
TEXT_EXTENSIONS = {
    ".py", ".txt", ".md", ".json", ".yaml", ".yml", ".toml", ".cfg", ".ini",
    ".js", ".ts", ".jsx", ".tsx", ".html", ".css", ".sql", ".sh", ".bash",
    ".java", ".c", ".cpp", ".h", ".hpp", ".go", ".rs", ".rb", ".php", ".xml",
    ".csv", ".log", ".env", ".gitignore", ".dockerfile", ".kt", ".swift",
    ".r", ".m", ".scala", ".lua", ".pl", ".ps1", ".bat", ".vue", ".svelte",
    ".rst", ".tex", ".ipynb", ".conf", ".properties", ".gradle", ".dart",
}

# Расширения, для которых заведомо не имеет смысла пытаться читать как текст
# (обычные бинарные форматы без текстового содержимого, не изображения и не
# документы со специальным разбором выше) — экономит время на эвристику.
KNOWN_BINARY_EXTENSIONS = {
    ".exe", ".dll", ".so", ".bin", ".dat", ".db", ".sqlite", ".sqlite3",
    ".mp3", ".mp4", ".avi", ".mov", ".mkv", ".wav", ".ogg", ".flac",
    ".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz", ".iso",
    ".ttf", ".otf", ".woff", ".woff2", ".pyc", ".class", ".jar", ".wasm",
}

MAX_TEXT_FILE_CHARS_IN_PROMPT = 6000    # столько символов файла отправляется в контекст LLM за раз
MAX_TEXT_FILE_CHARS_STORAGE = 200_000   # столько храним для полного просмотра файла в самом боте
MAX_IMAGES_PER_ARCHIVE = 6              # экономим токены/трафик — не более N картинок за раз
MAX_TEXT_FILES_PER_ARCHIVE = 15

# Расширения, по которым определяем формат архива (проверяются по имени файла,
# т.к. одно только содержимое не всегда однозначно отличает .tar.gz от .gz).
TAR_SUFFIXES = (".tar", ".tar.gz", ".tgz", ".tar.bz2", ".tbz2", ".tar.xz", ".txz")
ZIP_SUFFIXES = (".zip",)
PLAIN_GZIP_SUFFIXES = (".gz",)  # одиночный сжатый файл, НЕ tar.gz (проверяется после TAR_SUFFIXES)


class ZipSecurityError(Exception):
    """Архив не прошёл проверку безопасности (path traversal / архивная бомба / превышены лимиты).

    Название сохранено для обратной совместимости с уже существующим кодом бота
    (изначально поддерживался только .zip) — используется для ЛЮБОГО формата архива.
    """


@dataclass
class ExtractedImage:
    filename: str
    mime_type: str
    raw_bytes: bytes

    @property
    def data_base64(self) -> str:
        return base64.b64encode(self.raw_bytes).decode("ascii")

    def to_data_uri(self) -> str:
        return f"data:{self.mime_type};base64,{self.data_base64}"


@dataclass
class ExtractedTextFile:
    filename: str
    content: str                    # версия, урезанная под MAX_TEXT_FILE_CHARS_IN_PROMPT — идёт в контекст LLM
    full_content: str               # версия, урезанная только под MAX_TEXT_FILE_CHARS_STORAGE — для просмотра в боте
    truncated: bool = False         # True, если content короче full_content (урезан для экономии токенов LLM)
    storage_truncated: bool = False  # True, если даже full_content короче исходного файла (redko, очень большой файл)


@dataclass
class ExtractionResult:
    images: list[ExtractedImage] = field(default_factory=list)
    text_files: list[ExtractedTextFile] = field(default_factory=list)
    skipped_files: list[str] = field(default_factory=list)
    total_files_in_archive: int = 0

    @property
    def is_empty(self) -> bool:
        return not self.images and not self.text_files

    def summary(self) -> str:
        parts = [
            f"📦 Архив обработан: {self.total_files_in_archive} файлов всего.",
            f"🖼 Изображений для анализа: {len(self.images)}",
            f"📄 Текстовых/кодовых файлов: {len(self.text_files)}",
        ]
        if self.skipped_files:
            shown = self.skipped_files[:10]
            parts.append(
                "⏭ Пропущено (неподдерживаемый тип/слишком большой файл): "
                + ", ".join(shown)
                + (f" и ещё {len(self.skipped_files) - 10}" if len(self.skipped_files) > 10 else "")
            )
        return "\n".join(parts)


def detect_archive_format(filename: str) -> str | None:
    """Определяет формат архива по имени файла. Возвращает 'zip' / 'tar' / 'gzip' / None."""
    lower = filename.lower()
    for suf in TAR_SUFFIXES:
        if lower.endswith(suf):
            return "tar"
    for suf in ZIP_SUFFIXES:
        if lower.endswith(suf):
            return "zip"
    for suf in PLAIN_GZIP_SUFFIXES:
        if lower.endswith(suf):
            return "gzip"
    return None


def is_supported_archive(filename: str) -> bool:
    return detect_archive_format(filename) is not None


# ------------------------------------------------------------------ общие утилиты пути


def _safe_member_path(base_dir: Path, member_name: str) -> Path:
    """Строит путь к файлу внутри base_dir, отклоняя попытки выйти за её пределы
    (path traversal / zip-slip / tar-slip). Работает одинаково для zip и tar."""
    normalized = member_name.replace("\\", "/")
    if normalized.startswith("/") or (len(normalized) > 1 and normalized[1] == ":"):
        raise ZipSecurityError(f"Абсолютный путь в архиве запрещён: {member_name!r}")
    target = (base_dir / normalized).resolve()
    base_resolved = base_dir.resolve()
    if base_resolved != target and base_resolved not in target.parents:
        raise ZipSecurityError(f"Обнаружена попытка выхода за пределы архива (path traversal): {member_name!r}")
    return target


# ------------------------------------------------------------------ ZIP


def _validate_zip_archive(zf: zipfile.ZipFile) -> None:
    infos = zf.infolist()
    if len(infos) > MAX_FILES_IN_ARCHIVE:
        raise ZipSecurityError(
            f"В архиве слишком много файлов ({len(infos)} > {MAX_FILES_IN_ARCHIVE}). Отклонено."
        )

    total_uncompressed = 0
    for info in infos:
        if info.is_dir():
            continue
        if info.file_size > MAX_SINGLE_FILE_BYTES:
            raise ZipSecurityError(
                f"Файл {info.filename!r} слишком большой ({info.file_size} байт). Отклонено."
            )
        if info.compress_size > 0:
            ratio = info.file_size / max(info.compress_size, 1)
            if ratio > MAX_COMPRESSION_RATIO and info.file_size > 1024 * 1024:
                raise ZipSecurityError(
                    f"Файл {info.filename!r} имеет подозрительно высокую степень сжатия "
                    f"(похоже на архивную бомбу). Отклонено."
                )
        total_uncompressed += info.file_size
        if total_uncompressed > MAX_TOTAL_UNCOMPRESSED_BYTES:
            raise ZipSecurityError(
                f"Суммарный размер распакованных файлов превышает лимит "
                f"({MAX_TOTAL_UNCOMPRESSED_BYTES // (1024*1024)} MB). Отклонено."
            )


def _extract_zip(archive_path: Path, tmp_path: Path) -> None:
    try:
        with zipfile.ZipFile(archive_path) as zf:
            _validate_zip_archive(zf)
            for info in zf.infolist():
                if info.is_dir():
                    continue
                target_path = _safe_member_path(tmp_path, info.filename)
                target_path.parent.mkdir(parents=True, exist_ok=True)
                with zf.open(info) as src, open(target_path, "wb") as dst:
                    dst.write(src.read())
    except zipfile.BadZipFile as e:
        raise ZipSecurityError(f"Файл повреждён или не является zip-архивом: {e}") from e


# ------------------------------------------------------------------ TAR (.tar, .tar.gz, .tar.bz2, .tar.xz)


def _validate_tar_member(member: tarfile.TarInfo) -> None:
    if member.issym() or member.islnk():
        raise ZipSecurityError(
            f"Архив содержит символическую/жёсткую ссылку ({member.name!r}) — запрещено из соображений безопасности."
        )
    if member.isdev():
        raise ZipSecurityError(f"Архив содержит специальный файл устройства ({member.name!r}) — запрещено.")
    if member.size > MAX_SINGLE_FILE_BYTES:
        raise ZipSecurityError(f"Файл {member.name!r} слишком большой ({member.size} байт). Отклонено.")


def _extract_tar(archive_path: Path, tmp_path: Path) -> None:
    try:
        # mode="r:*" — автоопределение сжатия (gz/bz2/xz/без сжатия) по содержимому
        with tarfile.open(archive_path, mode="r:*") as tf:
            all_members = tf.getmembers()

            # Сначала проверяем АБСОЛЮТНО ВСЕ элементы архива (включая символические/
            # жёсткие ссылки и спецфайлы) — они должны быть отклонены явной ошибкой,
            # а не молча пропущены, чтобы у пользователя/админа не создавалось ложное
            # впечатление, что вредоносный архив был безопасно и полностью обработан.
            for member in all_members:
                _validate_tar_member(member)

            members = [m for m in all_members if m.isreg()]
            if len(members) > MAX_FILES_IN_ARCHIVE:
                raise ZipSecurityError(
                    f"В архиве слишком много файлов ({len(members)} > {MAX_FILES_IN_ARCHIVE}). Отклонено."
                )

            total_uncompressed = 0
            for member in members:
                total_uncompressed += member.size
                if total_uncompressed > MAX_TOTAL_UNCOMPRESSED_BYTES:
                    raise ZipSecurityError(
                        f"Суммарный размер распакованных файлов превышает лимит "
                        f"({MAX_TOTAL_UNCOMPRESSED_BYTES // (1024*1024)} MB). Отклонено."
                    )
                # Провалидируем путь заранее, чтобы не создавать файлы до полной проверки архива
                _safe_member_path(tmp_path, member.name)

            for member in members:
                target_path = _safe_member_path(tmp_path, member.name)
                target_path.parent.mkdir(parents=True, exist_ok=True)
                extracted = tf.extractfile(member)
                if extracted is None:
                    continue
                with open(target_path, "wb") as dst:
                    dst.write(extracted.read(MAX_SINGLE_FILE_BYTES + 1))
    except tarfile.TarError as e:
        raise ZipSecurityError(f"Файл повреждён или не является tar-архивом: {e}") from e


# ------------------------------------------------------------------ одиночный GZIP (не tar.gz)


def _extract_plain_gzip(archive_path: Path, tmp_path: Path, original_name: str) -> None:
    """Распаковывает одиночный .gz файл (например, script.py.gz), а не tar.gz архив."""
    inner_name = original_name[:-3] if original_name.lower().endswith(".gz") else "extracted_file"
    inner_name = Path(inner_name).name or "extracted_file"
    target_path = _safe_member_path(tmp_path, inner_name)

    try:
        with gzip.open(archive_path, "rb") as src:
            data = src.read(MAX_TOTAL_UNCOMPRESSED_BYTES + 1)
    except gzip.BadGzipFile as e:
        raise ZipSecurityError(f"Файл повреждён или не является gzip-файлом: {e}") from e
    except OSError as e:
        raise ZipSecurityError(f"Не удалось прочитать gzip-файл: {e}") from e

    if len(data) > MAX_TOTAL_UNCOMPRESSED_BYTES:
        raise ZipSecurityError(
            f"Распакованный файл превышает лимит размера "
            f"({MAX_TOTAL_UNCOMPRESSED_BYTES // (1024*1024)} MB) — похоже на архивную бомбу. Отклонено."
        )

    target_path.write_bytes(data)


# ------------------------------------------------------------------ единая точка входа


@contextmanager
def extract_archive_safely(archive_path: str | Path, original_filename: str | None = None) -> Iterator[Path]:
    """Контекстный менеджер: безопасно распаковывает архив (.zip, .tar[.gz|.bz2|.xz], .gz)
    во временную директорию и гарантированно удаляет её по выходу из блока `with`.

    `original_filename` используется для определения формата по расширению
    (если сам временный путь на диске не сохраняет исходное имя файла — так
    обычно и бывает при скачивании файла из Telegram).

    Использование:
        with extract_archive_safely(path, original_filename="project.tar.gz") as extracted_dir:
            ... работа с файлами внутри extracted_dir ...
        # тут временная директория уже удалена
    """
    archive_path = Path(archive_path)
    name_for_detection = original_filename or archive_path.name
    archive_format = detect_archive_format(name_for_detection)
    if archive_format is None:
        raise ZipSecurityError(
            f"Неподдерживаемый формат архива: {name_for_detection!r}. "
            f"Поддерживаются: .zip, .tar, .tar.gz/.tgz, .tar.bz2/.tbz2, .tar.xz/.txz, .gz"
        )

    with TemporaryDirectory(prefix="archive_extract_") as tmp_dir:
        tmp_path = Path(tmp_dir)
        if archive_format == "zip":
            _extract_zip(archive_path, tmp_path)
        elif archive_format == "tar":
            _extract_tar(archive_path, tmp_path)
        elif archive_format == "gzip":
            _extract_plain_gzip(archive_path, tmp_path, name_for_detection)
        yield tmp_path


def classify_and_read(extracted_dir: Path) -> ExtractionResult:
    """Проходит по распакованной директории, классифицирует файлы и готовит
    их содержимое (base64 для картинок, текст для кода/документов) для
    передачи в LLM. Поддерживает: изображения, обычный текст/код (по
    известным расширениям), PDF/DOCX/XLSX (через специальный разбор), а
    также ЛЮБОЙ ДРУГОЙ файл, который эвристически похож на текст (не входит
    в список заведомо бинарных расширений) — так что почти любой файл,
    оказавшийся внутри архива, будет так или иначе прочитан и передан модели.
    """
    result = ExtractionResult()
    all_files = sorted(p for p in extracted_dir.rglob("*") if p.is_file())
    result.total_files_in_archive = len(all_files)

    for path in all_files:
        rel_name = str(path.relative_to(extracted_dir))
        suffix = path.suffix.lower()

        if suffix in IMAGE_EXTENSIONS:
            if len(result.images) >= MAX_IMAGES_PER_ARCHIVE:
                result.skipped_files.append(rel_name + " (лимит картинок исчерпан)")
                continue
            try:
                raw = path.read_bytes()
                mime_type = mimetypes.guess_type(path.name)[0] or "image/png"
                result.images.append(
                    ExtractedImage(
                        filename=rel_name,
                        mime_type=mime_type,
                        raw_bytes=raw,
                    )
                )
            except Exception:  # noqa: BLE001
                result.skipped_files.append(rel_name + " (ошибка чтения)")
            continue

        if suffix in DOCUMENT_EXTENSIONS:
            if len(result.text_files) >= MAX_TEXT_FILES_PER_ARCHIVE:
                result.skipped_files.append(rel_name + " (лимит текстовых файлов исчерпан)")
                continue
            try:
                raw_text = extract_document_text(path, suffix)
            except Exception as e:  # noqa: BLE001
                result.skipped_files.append(f"{rel_name} (не удалось разобрать документ: {e})")
                continue
            result.text_files.append(_make_text_file(rel_name, raw_text))
            continue

        if suffix in KNOWN_BINARY_EXTENSIONS:
            result.skipped_files.append(rel_name + " (неподдерживаемый бинарный формат)")
            continue

        if len(result.text_files) >= MAX_TEXT_FILES_PER_ARCHIVE:
            result.skipped_files.append(rel_name + " (лимит текстовых файлов исчерпан)")
            continue
        try:
            raw_text = path.read_text(encoding="utf-8", errors="strict")
        except Exception:
            # Файл с неизвестным/нестандартным расширением, который не
            # получилось прочитать как чистый UTF-8 текст — скорее всего,
            # действительно бинарный формат, пропускаем его.
            if suffix in TEXT_EXTENSIONS or suffix == "":
                # Известное текстовое расширение — читаем с заменой
                # некорректных байтов вместо полного отказа.
                try:
                    raw_text = path.read_text(encoding="utf-8", errors="replace")
                except Exception:  # noqa: BLE001
                    result.skipped_files.append(rel_name + " (не удалось прочитать как текст)")
                    continue
            else:
                result.skipped_files.append(rel_name + " (не удалось прочитать как текст)")
                continue
        result.text_files.append(_make_text_file(rel_name, raw_text))

    return result


def _make_text_file(rel_name: str, raw_text: str) -> "ExtractedTextFile":
    """Обрезает текст под лимиты хранения/контекста и оборачивает в ExtractedTextFile."""
    storage_truncated = len(raw_text) > MAX_TEXT_FILE_CHARS_STORAGE
    full_content = raw_text[:MAX_TEXT_FILE_CHARS_STORAGE]
    truncated = len(raw_text) > MAX_TEXT_FILE_CHARS_IN_PROMPT
    prompt_content = raw_text[:MAX_TEXT_FILE_CHARS_IN_PROMPT]
    if truncated:
        prompt_content += "\n... [файл обрезан для контекста модели, полностью доступен в 📂 Файлах архива]"
    return ExtractedTextFile(
        filename=rel_name,
        content=prompt_content,
        full_content=full_content,
        truncated=truncated,
        storage_truncated=storage_truncated,
    )


# ------------------------------------------------------------------ извлечение текста из PDF/DOCX/XLSX


def extract_pdf_text(path: Path) -> str:
    """Извлекает текст из PDF постранично через pypdf. Возвращает пустую
    строку (не бросает исключение) для страниц без извлекаемого текста
    (например, отсканированные страницы-картинки без OCR-слоя) — вместо
    этого добавляет пометку, чтобы пользователь понимал, почему текста нет."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            raise ValueError("PDF защищён паролем, не удалось открыть без пароля")

    parts = []
    empty_pages = 0
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(f"--- Страница {i} ---\n{text.strip()}")
        else:
            empty_pages += 1
    if empty_pages and not parts:
        raise ValueError(
            "не удалось извлечь текст ни с одной страницы (вероятно, это скан "
            "без текстового слоя — OCR не поддерживается)"
        )
    if empty_pages:
        parts.append(f"\n[Примечание: {empty_pages} страниц не содержали извлекаемого текста]")
    return "\n\n".join(parts)


def extract_docx_text(path: Path) -> str:
    """Извлекает текст из .docx: абзацы по порядку плюс текст таблиц (ячейки
    через табуляцию, строки таблицы через перенос строки)."""
    import docx

    document = docx.Document(str(path))
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table_idx, table in enumerate(document.tables, start=1):
        table_lines = [f"\n--- Таблица {table_idx} ---"]
        for row in table.rows:
            table_lines.append("\t".join(cell.text.strip() for cell in row.cells))
        parts.append("\n".join(table_lines))
    if not parts:
        raise ValueError("документ не содержит извлекаемого текста")
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    """Извлекает содержимое .xlsx как текстовые таблицы: каждый лист —
    отдельный блок, строки через перенос, ячейки через ' | '. Лимитирует
    число листов и строк на лист, чтобы не раздувать контекст на огромных
    таблицах (см. MAX_XLSX_SHEETS/MAX_XLSX_ROWS_PER_SHEET)."""
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    sheet_names = wb.sheetnames[:MAX_XLSX_SHEETS]
    skipped_sheets = len(wb.sheetnames) - len(sheet_names)

    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        lines = [f"--- Лист: {sheet_name} ---"]
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= MAX_XLSX_ROWS_PER_SHEET:
                lines.append(f"[... остальные строки листа обрезаны, лимит {MAX_XLSX_ROWS_PER_SHEET} строк ...]")
                break
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                lines.append(" | ".join(cells))
            row_count += 1
        parts.append("\n".join(lines))

    if skipped_sheets > 0:
        parts.append(f"\n[Примечание: ещё {skipped_sheets} листов не показаны из-за лимита]")
    if not parts:
        raise ValueError("книга не содержит листов с данными")
    return "\n\n".join(parts)


def extract_document_text(path: Path, suffix: Optional[str] = None) -> str:
    """Единая точка входа для извлечения текста из документа со специальным
    разбором (PDF/DOCX/XLSX) по расширению файла."""
    suffix = (suffix or path.suffix.lower())
    if suffix in PDF_EXTENSIONS:
        return extract_pdf_text(path)
    if suffix in DOCX_EXTENSIONS:
        return extract_docx_text(path)
    if suffix in XLSX_EXTENSIONS:
        return extract_xlsx_text(path)
    raise ValueError(f"неизвестный тип документа: {suffix}")


def is_supported_document(filename: str) -> bool:
    """True, если файл — документ со специальным разбором (PDF/DOCX/XLSX),
    который умеет обрабатывать extract_document_text (в отличие от обычных
    текстовых/кодовых файлов, читаемых напрямую построчно)."""
    return Path(filename).suffix.lower() in DOCUMENT_EXTENSIONS


def read_single_document(path: Path, original_filename: str) -> "ExtractedTextFile":
    """Читает ОДИНОЧНЫЙ документ, присланный напрямую в чат (не внутри
    архива) — определяет способ разбора по расширению: специальный
    (PDF/DOCX/XLSX) или обычный текст/код (в т.ч. любой файл, который
    эвристически похож на текст)."""
    suffix = Path(original_filename).suffix.lower()

    if suffix in DOCUMENT_EXTENSIONS:
        raw_text = extract_document_text(path, suffix)
    elif suffix in KNOWN_BINARY_EXTENSIONS:
        raise ValueError("неподдерживаемый бинарный формат файла")
    else:
        try:
            raw_text = path.read_text(encoding="utf-8", errors="strict")
        except Exception:
            if suffix in TEXT_EXTENSIONS or suffix == "":
                raw_text = path.read_text(encoding="utf-8", errors="replace")
            else:
                raise ValueError("не удалось прочитать файл как текст")

    return _make_text_file(original_filename, raw_text)


def build_multimodal_user_content(
    user_text: str, extraction: ExtractionResult
) -> str | list[dict]:
    """Строит содержимое сообщения пользователя для OpenAI-совместимого API.

    Если в архиве есть картинки — возвращается мультимодальный список
    (текст + image_url блоки в формате, который понимает NVIDIA Vision API/NIM
    for VLMs — data URI с base64). Если картинок нет — обычная строка с текстом
    и вставленным содержимым текстовых файлов (что дешевле по токенам и
    работает с любой, даже не-vision, моделью).
    """
    text_parts = [user_text.strip()] if user_text.strip() else []

    if extraction.text_files:
        text_parts.append("\n📄 Содержимое файлов из архива:")
        for tf in extraction.text_files:
            note = " (обрезан)" if tf.truncated else ""
            text_parts.append(f"\n--- Файл: {tf.filename}{note} ---\n{tf.content}")

    if extraction.skipped_files:
        text_parts.append("\n" + extraction.summary())

    combined_text = "\n".join(text_parts) if text_parts else "Проанализируй содержимое архива."

    if not extraction.images:
        return combined_text

    content: list[dict] = [{"type": "text", "text": combined_text}]
    for img in extraction.images:
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": img.to_data_uri()},
            }
        )
    return content
